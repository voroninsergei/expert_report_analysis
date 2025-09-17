"""
Модуль для генерации DOCX‑отчёта на основе ответа ChatGPT.

Использует python‑docx для создания нового документа, добавления заголовка
и основного текста. Если необходимо, этот модуль можно расширить для
форматирования разделов, добавления таблиц или других элементов.
"""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document  # type: ignore
except ImportError as e:  # pragma: no cover
    Document = None  # type: ignore


def create_report(content: str, output_path: str | Path, title: str = "Сводный отчёт о недостатках") -> None:
    """Создаёт и сохраняет DOCX‑файл с переданным содержимым.

    :param content: текст, который будет записан в документ
    :param output_path: путь к файлу, который будет создан
    :param title: заголовок документа
    :raises ImportError: если python‑docx не установлен
    """
    if Document is None:  # pragma: no cover
        raise ImportError("python-docx не установлен. Добавьте его в зависимости.")

    doc = Document()
    # Добавляем заголовок
    doc.add_heading(title, level=1)
    # Добавляем основной текст. Используем splitlines, чтобы сохранить
    # структуру параграфов.
    for paragraph in content.splitlines():
        # Пропускаем пустые строки, чтобы добавить пустой параграф
        doc.add_paragraph(paragraph)

    # Сохраняем документ
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))